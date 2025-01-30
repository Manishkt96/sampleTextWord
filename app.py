import httpx
import os
import base64
import streamlit as st
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
import io
from PIL import Image

# Load environment variables
load_dotenv()

# Set API key for Gemini
gemini_flash = os.getenv("API_KEY")
gemini = genai.configure(api_key=gemini_flash)  # Set the API key globally


# Function to generate content with the images using the model
def generate_content(images, prompt):
    # Prepare the images and send to the model
    model = genai.GenerativeModel(model_name="gemini-1.5-pro")

    image_data = []
    for img in images:
        img_content = img.read()
        image_data.append({'mime_type': 'image/jpeg', 'data': base64.b64encode(img_content).decode('utf-8')})

    response = model.generate_content(image_data + [{'text': prompt}])
    return response.text


# Function to save extracted text to a Word document
def save_to_word(text):
    doc = Document()
    doc.add_heading('Extracted Text from Images', 0)
    doc.add_paragraph(text)

    # Save the document to a BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Streamlit UI
st.title("Image Object Detection")
st.write("Upload up to 5 images (JPG, PNG, GIF, BMP, TIFF) to identify the objects contained in the images.")

# Allow the user to upload up to 5 images
uploaded_files = st.file_uploader("Choose images", type=["jpeg", "jpg", "png", "gif", "bmp", "tiff"], accept_multiple_files=True)

# If images are uploaded
if uploaded_files:
    if len(uploaded_files) > 5:
        st.error("You can upload a maximum of 5 images.")
    else:
        # Process the images
        images = [file for file in uploaded_files]
        
        # Define the prompt
        prompt = "You are an OCR text extracter which extract text from image.Extract Text from images."
        
        # Process the images with the model
        if st.button("Convert"):
            response = generate_content(images, prompt)
            st.write("Objects detected in the images:")
            st.text(response)

            # Save extracted text to a Word document
            word_buffer = save_to_word(response)

            # Provide the user with a download link for the Word document
            st.download_button(
                label="Download Extracted Text as DOCX",
                data=word_buffer,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
